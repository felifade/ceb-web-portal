import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_DIR = "/Users/felipelopezsalazar/Library/Mobile Documents/com~apple~CloudDocs/School/CEB/01_Academico_y_Clases/Ciclo_2025-2026/Semestre_B_Feb-Jun/Material_2025-2026/Consulta Express/drive-download-20260701T205348Z-3-001"
OUTPUT_WORD = "/Users/felipelopezsalazar/Library/Mobile Documents/com~apple~CloudDocs/School/CEB/01_Academico_y_Clases/Ciclo_2025-2026/Semestre_B_Feb-Jun/Material_2025-2026/Consulta Express/Tabla_X-Press_Grupo_210.docx"

def main():
    # Obtener lista de archivos y extraer nombres
    docx_files = []
    for root, dirs, files in os.walk(INPUT_DIR):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                docx_files.append(file)
                
    if not docx_files:
        print(f"No se encontraron archivos en {INPUT_DIR}")
        return
        
    students = []
    for file in docx_files:
        name = file.replace("- Portafolio 2o. Parcial", "").replace(".docx", "").replace("_(1)", "").replace("_(2)", "").replace("_", "").strip()
        # Avoid duplicates since there are some (1) and (2) files
        if name not in students:
            students.append(name)
            
    students.sort()

    # Crear documento Word
    doc = Document()
    
    # Título principal
    title = doc.add_heading('Consulta X-Press – Participación estudiantil', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Grupo 210 | Evidencia de capturas de pantalla')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Espacio
    
    # Crear tabla con 2 columnas y tantas filas como alumnos + 1 (encabezado)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nombre del Alumno'
    hdr_cells[1].text = 'Captura de Pantalla'
    
    # Dar formato a los encabezados
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ajustar anchos (relativo)
    for row in table.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(4.5)
    
    # Llenar la tabla
    for student in students:
        row_cells = table.add_row().cells
        row_cells[0].text = student
        row_cells[1].text = "" # Espacio en blanco para la imagen
        
        # Darle algo de altura a la fila para que se vea el espacio para la imagen
        row_cells[1].paragraphs[0].paragraph_format.space_after = Pt(200)

    doc.save(OUTPUT_WORD)
    print(f"Documento Word generado exitosamente en: {OUTPUT_WORD}")

if __name__ == "__main__":
    main()
